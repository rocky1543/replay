# encoding: utf-8
import akshare as ak
import json
import logging
import re
import requests
import time
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from pyquery import PyQuery as pq

code_map = {}
zhang_ting_di_wei_tag = {}


def get_proxies():
    proxy_ip = get_proxy_ip()
    if proxy_ip == "":
        proxy_ip = get_proxy_ip()

    # 官网：https://www.qg.net/doc/1697.html
    authKey = "03WMRTUF"
    password = "8911F26173C9"
    print("proxy_ip:", proxy_ip)
    proxyUrl = "http://{}:{}@{}".format(authKey, password, proxy_ip)
    return {
        "http": proxyUrl,
        "https": proxyUrl
    }


def get_article_info(name):
    print("------------------------------")
    name = name.replace("Ａ", "A")
    proxies = get_proxies()
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"
    }
    print("proxies:", proxies)
    text = ""
    for _ in range(10):
        try:
            jiucai_url = 'https://www.jiuyangongshe.com/search/new?k={}股票异动解析'.format(name)
            print("jiucai_url:", jiucai_url)
            response = requests.get(jiucai_url, headers=headers, proxies=proxies, timeout=3)
            # print("response.text:", response.text)
            if response.status_code == 200 and response.text.count("股票异动解析") > 0:
                text = response.text

            # print("text:", text)
            time.sleep(0.2)
        except Exception as e:
            proxies = get_proxies()
            logging.error(e)
        if text:
            break

    href_pattern = re.compile(r'href="(/a/.*?)"')
    href_matches = href_pattern.findall(text)

    if len(href_matches) <= 0:
        return None

    # 输出匹配到的href
    for href in href_matches:
        for _ in range(10):
            try:
                href_full_url = "https://www.jiuyangongshe.com{}".format(href)
                print("href_full_url:", href_full_url)
                response = requests.get(href_full_url, headers=headers, proxies=proxies, timeout=3)
                text = ""
                if response.status_code == 200:
                    text = response.text

                doc = pq(text)
                title = doc(".fs28-bold")
                title = title.text()
                print("title:", title)
                print("name:", name)
                print("find:", title.find(name))
                if title.find(name) <= 0:
                    break

                ti_cai_text = doc(".mt40  > div.text-justify")
                ti_cai_text = ti_cai_text.text()
                print("ti_cai_text:", ti_cai_text)

                info = doc(".pre-line")
                date = get_today()
                try:
                    date = doc(".date").text().split(" ")[0]
                except Exception as e:
                    logging.exception(e)

                info = str(info)
                info = clean_preline_divs(info)
                code_info = code_map.get(name, None)
                tag = zhang_ting_di_wei_tag.get(name, "")
                if code_info:
                    change = code_info.get("涨跌幅", None)
                    code = code_info.get("代码", None)
                    print("change:", change)
                    print("code:", code)
                    if code and change and info:
                        info_arr = info.split("\n", 1)
                        info_0 = info_arr[0] + "  " + str(change) + "%" + "  " + tag + "  " + str(code)
                        info = info_0 + "\n" + info_arr[1]
                print("info:", info)
                return {"info": info, "date": date, "title": title, "ti_cai_text": ti_cai_text}
            except Exception as e:
                proxies = get_proxies()
                logging.exception(e)


def get_print_jiucai_url(name):
    print("------------------------------")
    jiucai_url = "https://www.jiuyangongshe.com/search/new?k=name&type=5".replace("name", name)
    print("jiucai_url:", jiucai_url)


def clean_preline_divs(html_string):
    """移除 pre-line div标签，保留内容"""
    pattern = r'<div class="pre-line" data-v-[a-zA-Z0-9]+="">(.*?)</div>'
    return re.sub(pattern, r'\1', html_string, flags=re.DOTALL)


def get_today():
    now = datetime.now()
    return now.strftime("%Y-%m-%d")


def save_word_text(bu_zhang_long, lian_ban, name_list, info_map, print_type="A5"):
    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 设置页面大小为A5
    section = doc.sections[0]
    if print_type == "A5":
        section.page_width = Cm(14.8)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # 题材
    add_page_break = False
    add_lian_ban_split = True
    add_he_xin_split = True
    for key in name_list:
        val = info_map.get(key, None)
        if not val or val is None:
            continue

        if print_type == "A5":
            # 分页符
            if not add_page_break:
                add_page_break = True
            else:
                doc.add_page_break()

        print("key:", key)
        print("val:", val)
        info = val.get("info")
        title = val.get("title")
        ti_cai_text = val.get("ti_cai_text")

        if key not in bu_zhang_long and add_lian_ban_split and add_he_xin_split:
            add_lian_ban_split = False
            doc.add_paragraph("- " * 90)

        if key not in lian_ban and key not in bu_zhang_long \
                and not add_lian_ban_split and add_he_xin_split:
            add_he_xin_split = False
            doc.add_paragraph("- " * 90)

        # 添加标题，0表示样式为title
        h1 = doc.add_heading(key, level=2)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加段落
        doc.add_paragraph(title + "\n" + info + "\n\n" + ti_cai_text + "\n")

    # 核心
    he_xin = """
*、觉悟的最大修为就是学会休息，盈利不是你抓住了很多机会，而是你放弃了更多的机会，学会give up吧
*、学会赚小钱，才能够赚大钱，空仓不是靠忍住，而是条件没形成，看不上
*、一个简单的高抛低吸游戏而已，为什么玩不赢呢，做好交易=过好生活而已
*、最折磨人的从来不是明确的坏，而是好得不够出众，一个月也就那一两次运气而已
*、第一天很重要：
    *、炒的【这段情绪】：风来了，情绪就来了，人就聚集了
    *、你要做的是：尽量在每个方向情绪开始的时候，第一天参与进去，尽量保证当天买点不亏钱
    *、市场在平静波动之间来回循环，当市场沉淀下来，波动小的时候，尽量找机会，波动大时风险大，尽量回避
    *、龙头低吸，二波，接力，全都靠子子孙孙回流，靠情绪弱转强；子子孙孙死的死，伤的伤，那就结束了
    *、每次情绪高潮的时候，就是悲剧的开始，已经没有形成向上情绪的条件了，习惯性低吸是亏钱的最大原因
    *、大盘不好，情绪不会太好，情绪溢价：赚的是第二天的情绪溢出效应，赚的是情绪转折节点的钱
    *、打板只有三种：首板，接力，二波，二波买的是启动点，龙头买的是连板中途的第一次大分歧或者再次二波    
*、大数定律
    *、赚多出的钱它也会亏回去，亏多出的钱它也会再赚回来，均值是会回归的
    *、去掉概率低的，保留概率高的，提升每一次出手的概率，当操作足够多的时候，盈利就会回归到均值上
    *、巨大的结果只需要日积月累的微小力量
    *、一样的条件产生一样的概率：
        *、时机：
            *、要看明天，明天可以持续，盈亏比要好
            *、恐慌，退潮，情绪低点，情绪弱转强，【这段情绪】开始的第一天低位，龙头确认给机会的第一天
        *、主线：大题材，大主流，大容量，梯队完整，主线清新，有高度，有宽度，万亿市场，3倍以上增长空间
        *、龙头：能带动市场，板块情绪，最猛，人气最旺，最吸金，是最核心的灵魂股，补涨他的子子孙孙很活跃
        *、二波：连板，横盘，辨识度，人气旺，天然合力，方向核心，量价完美，风韵犹存，刚刚开始的启动点
*、外重者内拙
    *、徒弟：师父，以我的资质，开悟需要多久？师父：十年
    *、徒弟：如果我加倍苦修呢？师父：二十年
    *、徒弟：如果我夜以继日，不休不眠呢？师父：那你将永无开悟之日
*、长短线：
    *、长线不做高位，短线不做低位，长线不做热门票，短线不做冷门票，长线看价值，中线看题材，短线看情绪
    *、长线只做市场极致绝望的低位，短线也一样，高度大龙头很久才会出来一个，情绪极致绝望的时候孕育大龙头
    *、长线只看价格高低，就像一年四季的气温，机会在市场极寒的时候，没有最好的公司，只有最好的价格
    *、短线只看情绪高低，就像海面的波浪，机会在波浪谷底，去做逆势抗争者，做市场弱转强
*、龙头：
    *、人气就是龙的魂，一切形态随着魂而动，人气可以通过观察子子孙孙活力而得
    *、龙头就是和指数、情绪起伏拟合的个股，是板块的脊梁
    *、板块主升，托举龙头主升，这才是是正解，社会讨论度决定板块地位，资金活跃度决定龙头地位
    *、人气是龙头的基石：
        *、虹吸效应：强的东西给机会，会吸干地位于他的标的的人气，资金，流动性
        *、束水攻沙：维持高标在高位强势需要大量的资金，资金需要收敛，凝聚到一个点
        *、人气 = 社会讨论度 + 资金活跃度，就好比居庸关的人气和八达岭就没法比
    *、相：
        *、叫她不凶，她心中已有凶相，而能做到和善，是对凶的克制，是戒
        *、而叫她和善，她心中已有和善之相，她和善只是顺其自然
*、光神：
    *、退潮的结果就是重新洗牌，推到重来，筛选掉一部分，保留一部分
    *、研究市场情绪的意义就在于如何让自己参与到每一波行情的起始
    *、再爆炸的行情，最多也就是三五天，螺旋式前进是必然规律，连续加速后就避开一下
    *、当最勇敢的资金有90%被割肉出来的时候，又构成了一个新的起点，开始新的游戏，周而复始
    *、在高潮中永远抱着敬畏的心态，在绝望中永远要充满希望，市场很差很差的时候，开始孕育新龙头
    *、我记起为何取一瞬流光这个id，当初的本意，只是吃情绪转好的那一瞬，氛围转强的时候去做人气最高的个股
    *、三点共振：1、指数大势，单边上涨；2、当日核心板块，随指数同步启动；3、个股为板块内人气容量核心
    *、只有类似东财这种级别的东西不断出现，好的赚钱效应才能维持下去，所以不用预判以后会不会好
    *、你就做好每一阶段的龙头就行了，把市场简化成一波一波的，做好每一个龙头实际上你就完整参与了这一波牛
*、92科比 & 群总：
    *、要么最高，要么最低，最高是龙头，最低是补涨，切换，人气在哪就去哪儿找机会
    *、接力的核心就是人气，龙头自身要有人气，板块要有人气，子子孙孙要有活力
    *、高位不行了，就去做低位的票，从低位从新做起来
    *、你能团结多少人，你就能干多少事，人气所在，牛股所在，流动性产生溢价
    *、研究这研究那有个毛用，要么空仓，要么核心
*、杂文：
    *、就跟打游戏一样，别人的技能没放完，你就贸然出击不就是找死嘛
    *、人的一切执念都是源自于对有的追求，有求于外者内拙
    *、就像小时候养的小鸡小鸟，过度的关心它，反而会把它弄死
    *、搞懂解决什么具体问题，影响多少人，影响多少钱，为什么小米值那么多钱？是因为智能手机市场值那么多钱
    *、为什么卖了之后，它跌你就高兴，它涨你就难过呢
*、方法论：
    *、做【首板】就是看高做低，高位拉出空间，【龙头】就是做二板之上谁最有人气，低位子子孙孙会奶他
    *、龙头：班级学霸；补涨龙=大题材里，高低切；补涨=大题材里，看高做低;补涨=跟风=蹭热度
老师：
    *、复盘看看当天炒什么题材，再看看有没有可能二波的
    *、盘中看看今天炒什么题材，看看有没有【首板】可能二波的
    *、老师主要做的还是二波，找对方向，做二波，图已经做好了，资金已经准备好了，就等方向切过来
    *、大盘暴跌就是好机会，高位只做人气龙头，其他都是去低位找启动点
养家：
    *、观察赚钱效应：观察人气回流的过程，观察亏钱效应：观察人气消散的过程
    *、整个循环：低位的赚钱，慢慢延伸到高位赚钱，然后高位出现亏钱，再回到低位，高低位都亏钱
*、指数：
    *、作用是预判大环境【亏钱效应】和【赚钱效应】，市场不好，题材也走不高
    *、高位靠低位奶，大盘推板块向上，板块推龙头向上，这种向上的力是最强的
情绪周期：
    *、核心思想【他带起的情绪】，会有很多子子孙孙追随他，奶他；九阳神功：动于阴末，止于阳极
    *、在大龙头倒下的时候，退潮强大的亏钱效应，会将带有"高位"属性的票冲得溃不成军
    *、一个板块核心人气股涨得猛，代表板块情绪好，核心人气都不行，代表情绪不好
        *、大部分的股票都是随波逐流的，去看他们也没有什么意义，只有那些有辨识度的股票，
        *、才能够真正的吸引资金，他们的强弱，能够带动板块走强或走弱
        *、这个时候你就知道市场情绪是好是坏，所以这些辨识度个股，其实是整个市场情绪的指示灯
    *、大盘环境好的时候，会助长龙头情绪，大盘不好的时候，会抑制龙头情绪
目标：
    *、资金总是往前排切，往低位切，所以前排和低位才是安全的，所以目标要么前排核心，要么低位首板
    *、你关注的颗粒度应该是每一个交易，一定要删除，简化到可以落地，如果条件不收敛，就永远总结不出规律
    """
    doc.add_page_break()
    doc.add_paragraph(he_xin)

    save_dui(doc)

    # 方法三：同时设置宽高（保持比例）
    doc.add_picture('./市场资金潮汐.jpg', width=Inches(4.5))
    # doc.add_picture('./市场资金潮汐.jpg', width=Inches(3), height=Inches(2))
    # doc.add_picture('./市场资金潮汐.jpg')

    doc.save('result/复盘.docx')

    save_tj(doc)
    save_js(doc)
    save_gz(doc)


def save_sc(doc):
    # delete_paragraph(doc)
    doc.add_page_break()
    # 条件
    condition = ""
    for line in open("./市场2部分.txt").readlines():
        condition = condition + line

    doc.add_paragraph(condition)
    # doc.save('result/条件.docx')


def save_dui(doc):
    # delete_paragraph(doc)
    doc.add_page_break()
    # 条件
    condition = ""
    for line in open("每天计划.txt").readlines():
        condition = condition + line

    doc.add_paragraph(condition)
    # doc.save('result/条件.docx')


def save_tj(doc):
    delete_paragraph(doc)
    # 条件
    condition = ""
    for line in open("总结/条件.txt").readlines():
        condition = condition + line

    doc.add_paragraph(condition)
    doc.save('result/条件.docx')


def save_js(doc):
    delete_paragraph(doc)
    # 龙头断板计数
    condition = ""
    for line in open("总结/龙头断板计数.txt").readlines():
        condition = condition + line

    doc.add_paragraph(condition)
    doc.save('result/龙头断板计数.docx')


def save_gz(doc):
    delete_paragraph(doc)
    # 共振
    condition = ""
    for line in open("总结/共振.txt").readlines():
        if line.strip().startswith("回归常识"):
            doc.add_paragraph(condition)
            doc.add_page_break()
            condition = ""
        condition = condition + line

    doc.add_paragraph(condition)
    doc.save('result/共振.docx')


def delete_paragraph(doc):
    for paragraph in list(doc.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


def get_timestamp(date):
    from dateutil import parser
    return int(parser.parse(date).timestamp())


def get_zhang_ting_list(file):
    name_list = []
    for line in open(file):
        if line.strip():
            data = line.strip().split(",")
            name_list.append(data[0])
            if len(data) == 2 and data[1]:
                zhang_ting_di_wei_tag[data[0]] = data[1]
    return name_list


def get_zhang_ting_map(ti_cai_list):
    zhang_ting_map = {}
    for val in ti_cai_list:
        zhang_ting_map[val] = get_zhang_ting_list("./input/{}.txt".format(val))

    return zhang_ting_map


def get_code_map():
    # 获取东方财富网-沪深京 A 股-实时行情
    df = ak.stock_zh_a_spot_em()
    for index, row in df.iterrows():
        code_map[row["名称"].strip()] = {"代码": row["代码"].strip(), "涨跌幅": row["涨跌幅"]}


def get_proxy_ip():
    try:
        # 青果网络的API地址和参数：https://www.qg.net/tools/IPdebug.html
        api_url = "https://share.proxy.qg.net/get?key=03WMRTUF&num=1&distinct=true"

        response = requests.get(api_url)
        data = json.loads(response.text)
        data = data.get("data", [])

        server_list = [val.get("server") for val in data]
        if server_list and len(server_list) > 0:
            return server_list[0]
    except Exception as e:
        logging.exception(e)
    return ""


def get_name_list(file, filter_list):
    name_list = []
    for line in open(file):
        line = line.strip()
        data = line.strip().split(",")
        name = data[0].strip()

        if len(data) == 2 and name:
            tag = data[1].strip()
            zhang_ting_di_wei_tag[name] = tag

        if name and name not in filter_list:
            name_list.append(name)

    return name_list


def get_replay_name_list():
    bu_zhang_long = get_name_list("input/1-龙头属性.txt", [])
    lian_ban = get_name_list("input/2-连板情绪.txt", bu_zhang_long)
    filter_list = bu_zhang_long + lian_ban
    fu_pan = get_name_list("input/3-大盘核心.txt", filter_list)
    print("bu_zhang_long:", bu_zhang_long)
    print("lian_ban:", lian_ban)
    print("fu_pan:", fu_pan)

    return bu_zhang_long, lian_ban, bu_zhang_long + lian_ban + fu_pan


if __name__ == '__main__':

    # 获取个股代码
    # get_code_map()

    # 获取题材涨停的个股
    bu_zhang_long, lian_ban, name_list = get_replay_name_list()
    print("name_list:", name_list)

    # 爬取涨停数据
    info_map = {}
    for name in name_list:
        # get_print_jiucai_url(name)
        # name_id = input("股票名字id：")
        # print("name_id:", name_id)
        # article_info = get_article_info_v2(name_id)
        article_info = get_article_info(name)
        if article_info:
            info_map[name] = article_info

    print("info_map:", info_map)
    # 保存到word
    save_word_text(bu_zhang_long, lian_ban, name_list, info_map, "A4")
