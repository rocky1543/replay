# encoding: utf-8
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt, RGBColor


def save_word_text(print_type="A5"):
    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 设置页面大小为A5
    section = doc.sections[0]
    if print_type == "A5":
        section.page_width = Cm(14.8)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    h1 = doc.add_heading("复盘", level=2)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fu_ban = "复盘三步曲：\n" \
             "1、分析主流板块: 大的主流在哪里，主流没走完，就只能在这个主流里面选股\n" \
             "2、分析时机：通过观察连板天梯的强弱转换，分析情绪强弱变化\n" \
             "3、分析它领涨的个股：寻找主流中各个小周期领涨股\n"
    doc.add_paragraph(fu_ban)
    doc.add_paragraph("- " * 90)

    # 添加标题，0表示样式为title
    yu_lu_list = [
        "要么做绝对的高位，要么做绝对的低位，要么点带面，要么面带点",
        "纯做龙头主要看高标就行，能带动板块上涨的才叫龙头，高标只能活一个，除非有一个太强买不到，半生其他龙头，"
        "龙头就是特别强的个股，贯穿整个题材行情炒作，板块指数跟着它走，也就是它带着板块其他个股在走",
        "模式应该是去找人气第一名，有中心思想的总结规律",
        
        "龙头高位断板大分歧会耗尽力量，跌停直接宣布退潮死掉，退潮了4天，市场底部开始变好，满足了市场共振的条件，这个时候出现一个大题材，"
        "满足了题材共振的条件，其次再去找满足龙头的条件，然后硬怼，此时的龙头，大概率是新题材的最高标，因为老周期的标的大部分都在调整，"
        "缺少题材共振的条件，龙头确认的那天，就是新周期开始的第一天，经历1,2,3,4天之后，大概率龙头重复断板，然后经历1,2,3,4天之后，"
        "又开始一个新周期，龙头和他脚下的小弟倒下的时候，会不会产生巨大的振动，新东西想活，就看它能不能顶得住龙头死的退潮考验",

        "核心研究的问题：这情绪还能不能延续",
        "永远只研究拐点，只研究高低，中间过程不必考虑，太阳升起的时候，就不要只思考自己那点事了，"
        "太阳就是市场最有人气，最靓的东西",
        "知道哪些是核心之后，就等啊，等自己的机会出现",
        "龙头做出来的人都是将帅之才，是万里挑一，非常勇猛的人，是目标明确，不拘泥小节，只做大事的人，"
        "贪图小利者，容易吃大亏，既然要做大事，就要知道什么是大事",
        "题材和龙头是线性关系，题材多大，龙头就多大，买点 = 大题材 + 大龙头 + 好时机 + 好位置",

        "世间万事万物都是诸多的因缘条件和合而成，是运动的，发展的，变化的，缘起而生，缘散而灭，而非永恒固定在某种状态的，"
        "不以人的意志为转移，世人却执着于事物停留在某种状态，从而产生喜怒哀乐，悲欢离合的情绪，你看这火是不是很旺，"
        "但过一会它就过去了，美好成为美好的时候，它具备的条件已经开始变化了，它就开始慢慢消亡，像这火烧得很旺的时候已经在慢慢燃尽了",

        "为什么有的人越努力越焦虑？因为他们只看到了努力，却忽视了规律，就像老百姓种地，再着急也要遵循春种秋收"
        "的自然规律，规律是在一定条件下能够稳定、重复发生的事情",
        "看盘：先大后小，自上而下，关注主要，忽略次要，人性之上，规律之下",
        "龙头就是去理解那个在充分条件时产生的砸不死的合力，如果砸死了，证明买的票是垃圾，割肉就好了",
        "赚钱的本质一定是低买高卖：1、情绪足够低位，2、向上有足够强的动力",
        "思考时不能加入太多的变量，不然复杂度会变得非常高，脑子就会混沌",
        "市场没有量，其实就是玩的人少了，条件不充分时的大胆，是你亏钱的根源",
        "在市场极其弱的时候，情绪转强可能很快就到来，这时去找和大盘，板块共振向上的最强个股，"
        "你为什么做不好，因为符合条件的机会是极其少，你却又有一颗无限想操作的心",
        "如果你在顶层做了正确的事，底层结果就不会差，一定要知道市场好绕不开哪个板块，板块好绕不开哪个个股",

        "心不死，道不生：好坏，美丑，穷富都是矛盾对立的双方，如果你不能放下其中一方，矛盾双方就在你心里斗争，"
        "内耗，折磨你，直至你不在乎它，它才会消失",

        "",
        "抓住主要：",
        "要善于抓大放小，主要问题，重大决策一定要做好，大大小小的事都重视，必然会精力分散，影响重大决策",
        "敌强我弱是客观存在的问题，如何保存力量是主要的矛盾，做好战略防御，及时退却，使自己立于主动地位",

        "其实不需要智慧，只需要做好选择，行情不好的时候，不要出手，行情好的时候，猛干核心",
        "二波本质是让你看清楚个股的地位，辨识度，人气，群众基础，看清楚谁是主要个股，看清楚后不等于要买，"
        "厉害的股票肯定有二波，大鹏展翅，肯定要两个翅膀，第一个翅膀就是行情的第一波，咱们不管他，咱们就盯着第二个翅膀的出现，"
        "盯着可能有二波的个股，在从其中找规律，不断的练，不断的亏钱，信心一点点的被打碎，又一点点的捡起来",
        "要从全局考虑，从高低位、新旧题材四个纬度，找出关键点在哪里，明确该做什么，不该做什么",

        "",
        "绝对优势：",
        "当局势有优势，辅助都很强，大哥就可以在团战中三杀，四杀，暴走，反之酱油秒躺，大哥也站不住",
        "打仗没有什么神秘，打得赢就打，打不赢就走，你打你的，我打我的",
        "任何一场博弈，都是以多胜少，以强胜弱，牛刀杀鸡，狮子打兔子，要有绝对的，压倒性的优势，"
        "如果没有这个条件只能等，在绝对实力面前，一切技术都是花里胡哨",

        "",
        "退潮:",
        "市场大哥都死了，本质是连板天梯(高位股)撑不住了，其他大哥也不能活，板块大哥都死了，就不要去板块里面玩了",
        "开始分歧了，后排都要卖，当龙头也断板了，不涨了，最漂亮的那个也没人追了，说明买盘力量已经枯竭了，"
        "高位已经撑不住了",

        "最高标断板的时候，第1、2天要把高位股全卖了，第2、3天情绪会惯性从顶部往下杀，第4天之后"
        "大部分的中高位都死的差不多了，下杀的情绪释放差不多了，高标可能只剩下1、2个，目标少了就清晰了，"
        "也就是最高标断板，会带着好几个高标跟着它一起死，第4天之后，等底部情绪好了，底部会反哺，反推高标，"
        "能不能做，看情绪转强的力度",

        "为什么大杀之后老东西就不行了，因为个股中会有很多获利、埋伏、套牢盘，筹码很复杂，筹码已经坏掉了",
        "怎么判断题材有没有结束：龙头核心A杀，当下属于退潮阶段，龙头人气没散，板块指数量能很足，题材炒作就没结束，"
        "板块小弟很活跃，龙头核心就不会差",

        "",
        "其他：",
        "要买就买那个最独一无二的，最强的，宁可不做也要缩小范围，减少目标，你敢在这个垃圾上冒险，"
        "就敢在其他类似的垃圾上冒险",
        "市场不认可的东西都是垃圾，如果从个股看不出确定性，多看看板块，确定性就有了",
        "把很复杂的东西，简化成你能操作的东西，考虑的东西越少越好，现在最强的东西是什么，用眼睛看就可以了，不用去想",

        "华映科技牛逼吗？在我眼里就是个垃圾，他涨一万个板也是垃圾，常山，海能达，四川这些才是大道",

        "主升浪，二波，其实就是做首板，这种首板往往都是发生在情绪好的时候，因为大资金不会瞎发动，而情绪的好坏都会反应在高标身上",
        "夫战，勇气也。一鼓作气，再而衰，三而竭，彼竭我盈，故克之，夫大国，难测也，惧有伏焉，吾视其辙乱，望其旗靡，故逐之",
        "这几年一路走来的艰难困苦，有什么值得歌颂的呢，本质上不是一个有天赋的人，打游戏如此，打球如此，智商也如此",
        """解决问题的步骤：
            1、万能公式：事物变化的底层规律 = 事物固有本性 + 外在条件
            2、从全局的角度，在纷繁复杂的问题中找出主要问题，也就是找出主要矛盾
            3、分析怎么样才能解决这个问题，也就分析事物固有性质，特点是什么，需要提供的外在条件是什么
            4、拆分条件，哪个是主要，哪个是次要，以主要条件为全局，按1，2步骤递归分析处理：
                1、内因，即事物固有性质
                2、外在条件
            
        """
    ]
    zhu = ""
    line = 0
    for yu_lu in yu_lu_list:
        if yu_lu:
            zhu = zhu + "{}、{}\n".format(line + 10, yu_lu)
            line = line + 1
        else:
            zhu = zhu + "\n"
    # 添加段落
    doc.add_paragraph(zhu)
    # doc.add_page_break()

    # 保存文档
    bj = """关注的主要三大范围：时间节点、题材方向、核心个股
    一、时间节点：涨潮， 退潮，随着时间变化，买卖强弱关系是会相互转化的，即阴阳循环
    二、代码简写：
        if "时间节点" == "最高标断板～大杀彻底之前":
            "尽可能放弃"
        else:
            if "方向" == "主线题材" and "个股" == "核心龙头":
                "可以考虑"
            else:
                "尽可能放弃"
    """

    gui_lv = """规律是事物存在某种本性，且在一定条件下可以稳定，重复发生的事，如果事物没有那种本性，无论提供什么条件，都不会产生相应的变化  
    一、万能公式：事物变化的底层规律 = 事物固有本性 + 外在条件
    二、固有本性：
        1、题材三大原则：大新强，大主流，大未来，大容量
        2、龙头：
            1、有资金进场维护，有完美的量价配合，放量突破新高，有人气
            2、没有透支行情，分歧整体向上
            3、形成好位置：
                1、调整充分后的二波首板，N型首版
                2、趋势主升浪，强中强原则：创新高，强于自身历史，强于同行，强于板块，强于大盘
                3、逆势大盘，强于大盘
            4、容量核心，换手充分
            5、逻辑纯正
            6、有辨识度，市场最独一无二的标的    
        3、事物固有本性例子：鸡蛋能变成小鸡，石头不能
    
    三、外在条件(通过增加条件，减少操作)：
        1、必要条件：
            1、情绪最好足够低位，最好是连续大杀几天，杀了再杀之后的低点，和跳绳一样，绳子落地了再参与 
            2、大盘必须转强，放量向上
            3、板块指数要强于大盘，强于其他板块，板块中多个个股走出新高状态
            4、板块中核心个股要强于板块
            5、情绪转强，和大盘，板块，龙头，大资金共振向上
            6、资金主动，买盘要有强劲的动力，要有压倒性优势
            7、板块氛围好，小弟很多很强
            
        2、加分条件：
            1、消息刺激，新闻发酵，吸引人气，带动情绪
            2、情绪转强竞价：板块高标百万封单顶一字，容量核心竞价抢筹
            
        3、不利条件：坚决不做
            1、龙头断板，连板天梯中的中高位股撑不住，高位整体向下杀，要学会放弃
            2、大盘有系统性风险，暴跌带崩情绪，要学会放弃   
              
        4、外在条件例如：鸡蛋需要特定的温度、湿度才能成小鸡（其实需要的条件很简单）
        
    万能公式：           
        事物变化的底层规律 = 事物固有本性 + 外在条件
        本性：本身固有的性质
            1、阳性质：规定事物的现在，直接显露，例如白纸的颜色是白的
            2、阴性质：规定事物的未来，隐藏在体内，想显露出来需要外在条件刺激，例如纸遇火可燃
        条件：1、充要条件，2、充分条件，3、必要条件   
    """
    doc.add_paragraph(bj)
    # doc.add_paragraph(gui_lv)
    doc.save('result/规律.docx')


if __name__ == '__main__':
    # 保存到word
    save_word_text("A4")
