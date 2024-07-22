# encoding: utf-8
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt, RGBColor

keyword_map = {}


def save_word_text():
    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 设置页面大小为A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    fin = open("./keyword_text")
    count = 1
    for line in fin.readlines():
        line = line.strip()
        if not line:
            continue
        if line.count("**") <= 0:
            continue
        line = line.replace("**", "")
        print("line:", line)
        line = line.split(".", 1)[1]

        line_data = line.split("：", 1)
        print("line_data:", line_data)
        keyword = line_data[0]
        if keyword_map.get(keyword):
            continue
        keyword_map[keyword] = 1

        # 添加段落
        p = doc.add_paragraph()

        # 段落中的关键字
        run = p.add_run("{}. {}：".format(count, keyword))
        run.bold = True
        # 字体倾斜：
        # run.italic = True
        # 下划线：
        # run.underline = True
        # 字体颜色：
        # run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        # 字体大小：
        # run.font.size = Pt(12)  # 12磅

        # 段落中的关键字解释
        p.add_run(line_data[1])
        count = count + 1

    # 保存文档
    doc.save('关键字.docx')


if __name__ == '__main__':
    save_word_text()
