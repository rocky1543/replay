# encoding: utf-8
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt, RGBColor


def save_word_text(lian_ban_data, print_type="A5"):
    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 设置页面大小为A5
    section = doc.sections[0]
    if print_type == "A5":
        section.page_width = Cm(14.8)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # 添加标题，0表示样式为title
    h1 = doc.add_heading("连板天梯", level=2)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("物壮则老：\n"
                      "1、龙头断板第4天，老周期的应该就全死完了, 就可以博弈新周期的最高标了\n"
                      "2、目的：观察主流板块中各个小周期情绪的起起伏伏，强弱转换，情绪转强的时候进，高潮的时候出")
    for line in lian_ban_data:
        line = line.replace("\t", ":\n")
        line = line.replace("&", "\n")
        # 添加段落
        doc.add_paragraph(line)

    doc.save('result/连板天梯.docx')


def get_day():
    # 获取当前时间
    now = datetime.now()
    return now.strftime("%Y%m%d")


def get_lian_ban_data():
    lian_ban_map = {}
    fo = open("./input/复盘核心.txt", "w")
    for line in open("./input/连板天梯.txt"):
        data = line.split("\t")
        name = data[0].strip()
        ban_num = data[-1].strip()
        fo.write(name + "," + ban_num + "板\n")

        lian_ban_list = lian_ban_map.get(ban_num, [])
        lian_ban_list.append(name)
        lian_ban_map[ban_num] = lian_ban_list

    data = ""
    for key, val in lian_ban_map.items():
        data = data + "{}板: {}".format(key, ", ".join(val)) + "&"

    day = get_day()
    data = day + "\t" + data
    print("data:", data)

    all_data = []
    day_list = []
    for line in open("./input/连板天梯_history.txt"):
        bef_day = line.strip().split("\t")[0]
        if bef_day in day_list:
            continue
        day_list.append(bef_day)
        all_data.append(line.strip())

    print("day:", day)
    print("day_list:", day_list)
    if str(day) not in day_list:
        all_data.insert(0, data)

    fo = open("./input/连板天梯_history.txt", "w")
    for line in all_data[:6]:
        print("all_data:", line)
        fo.write(line.strip() + "\n")

    return all_data


if __name__ == '__main__':
    lian_ban_data = get_lian_ban_data()
    save_word_text(lian_ban_data, "A4")
